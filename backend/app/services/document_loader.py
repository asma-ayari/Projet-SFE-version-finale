"""
Service de chargement et découpage des documents.
Supporte TXT, PDF, DOCX.
Découpage intelligent : par article, par cours, ou par section arabe.
"""
import re
from pathlib import Path
from typing import List
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
from app.core.config import settings


# ---------------------------------------------------------------------------
# Loaders
# ---------------------------------------------------------------------------

def _load_txt(file_path: Path) -> str:
    return file_path.read_text(encoding="utf-8")


def _load_pdf(file_path: Path) -> str:
    from PyPDF2 import PdfReader
    reader = PdfReader(str(file_path))
    text_parts = []
    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            text_parts.append(page_text)
    text = "\n\n".join(text_parts)

    # Alerte qualité extraction PDF: certains PDF (notamment arabe scanné)
    # produisent un texte corrompu difficilement exploitable par le RAG.
    replacement_count = text.count("�")
    if replacement_count >= 20:
        print(
            f"⚠️ Extraction PDF potentiellement corrompue: {file_path.name} "
            f"(caractères invalides détectés: {replacement_count})."
        )
        print(
            "   👉 Recommandé: convertir ce PDF en .txt (OCR si nécessaire) "
            "et ré-ingérer les documents."
        )

    return text


def _load_docx(file_path: Path) -> str:
    import docx
    doc = docx.Document(str(file_path))
    return "\n\n".join(para.text for para in doc.paragraphs if para.text.strip())


FILE_LOADERS = {
    ".txt": _load_txt,
    ".pdf": _load_pdf,
    ".docx": _load_docx,
}


# ---------------------------------------------------------------------------
# Chargement
# ---------------------------------------------------------------------------

def load_documents() -> List[Document]:
    documents = []
    docs_dir = Path(settings.DOCUMENTS_DIR)

    if not docs_dir.exists():
        print(f"⚠️ Dossier {docs_dir} introuvable. Création...")
        docs_dir.mkdir(parents=True, exist_ok=True)
        return documents

    for lang_dir in docs_dir.iterdir():
        if lang_dir.is_dir() and lang_dir.name in ("fr", "ar"):
            language = lang_dir.name
            for file_path in lang_dir.iterdir():
                if file_path.suffix.lower() in FILE_LOADERS:
                    loader = FILE_LOADERS[file_path.suffix.lower()]
                    try:
                        print(f"📄 Chargement : {file_path.name} ({language})")
                        content = loader(file_path)
                        if content.strip():
                            documents.append(Document(
                                page_content=content,
                                metadata={
                                    "source": file_path.name,
                                    "language": language,
                                    "file_type": file_path.suffix.lower(),
                                }
                            ))
                        else:
                            print(f"   ⚠️ Fichier vide, ignoré")
                    except Exception as e:
                        print(f"   ❌ Erreur: {e}")

    print(f"✅ {len(documents)} document(s) chargé(s)")
    return documents


# ---------------------------------------------------------------------------
# Découpage intelligent
# ---------------------------------------------------------------------------

def split_documents(documents: List[Document]) -> List[Document]:
    """
    Choisit automatiquement la meilleure stratégie de découpage selon le document.
    """
    all_chunks = []

    for doc in documents:
        chunks = _smart_split(doc)
        all_chunks.extend(chunks)
        print(f"   → {doc.metadata['source']} : {len(chunks)} chunks")

    print(f"✅ {len(all_chunks)} chunk(s) total créés")
    return all_chunks


def _get_custom_pattern(doc: Document) -> str:
    """Retourne un pattern regex personnalisé selon la langue du document."""
    language = doc.metadata.get("language")
    if language == "fr":
        return settings.CUSTOM_SPLIT_PATTERN_FR.strip()
    if language == "ar":
        return settings.CUSTOM_SPLIT_PATTERN_AR.strip()
    return ""


def _smart_split(doc: Document) -> List[Document]:
    """Détecte le type de document et applique la bonne stratégie."""
    content = doc.page_content

    if settings.CHUNKING_MODE.lower() == "recursive":
        print("   ℹ️ Mode recursive activé (découpage par taille)")
        return _fallback_split(doc)

    # 0. Pattern personnalisé (prioritaire)
    custom_pattern = _get_custom_pattern(doc)
    if custom_pattern:
        try:
            chunks = _split_by_pattern(
                doc,
                pattern=custom_pattern,
                label="pattern personnalisé"
            )
            if chunks:
                return chunks
            print("   ⚠️ Pattern personnalisé non matché, fallback smart")
        except re.error as e:
            print(f"   ⚠️ Regex personnalisée invalide: {e}")

    # 1. Articles français : "Article 1 :", "Article 12 :"
    if re.search(r"Article\s+\d+\s*:", content, re.IGNORECASE):
        chunks = _split_by_pattern(
            doc,
            pattern=r"(?=^Article\s+\d+\s*:)",
            label="article FR"
        )
        if chunks:
            return chunks

    # 2. Chapitres français : "=== CHAPITRE X ==="
    if re.search(r"===.+===", content):
        chunks = _split_by_pattern(
            doc,
            pattern=r"(?=^===\s*.+\s*===)",
            label="chapitre FR"
        )
        if chunks:
            return chunks

    # 2.b Titres FR génériques : CHAPITRE/SECTION/TITRE/BLOC en début de ligne
    if re.search(r"^(?:CHAPITRE|Chapitre|SECTION|Section|TITRE|Titre|PARTIE|Partie)\b", content, re.MULTILINE):
        chunks = _split_by_pattern(
            doc,
            pattern=r"(?=^(?:CHAPITRE|Chapitre|SECTION|Section|TITRE|Titre|PARTIE|Partie)\b.*$)",
            label="titres FR",
            flags=re.MULTILINE
        )
        if chunks:
            return chunks

    # 3. Cours pédagogiques PDF : "1er cours :", "2ème cours :", "Cours X :"
    if re.search(r"(cours\s*\d*\s*:|◼\s*\d)", content, re.IGNORECASE):
        chunks = _split_by_pattern(
            doc,
            pattern=r"(?=(?:◼\s*\d|(?:Le\s+)?\d+(?:er|ème|e)?\s+cours\s*:|Cours\s+\w+\s*:))",
            label="cours pédagogique",
            flags=re.IGNORECASE
        )
        if chunks:
            return chunks

    # 3.b Supports FR pagines : "Page 3/6" + sections pedagogiques
    if re.search(r"(?:page\s*\d+\s*/\s*\d+|\b(?:protection|alerte|secours|signalisation|priorit[ée])\b)", content, re.IGNORECASE):
        chunks = _split_by_pattern(
            doc,
            pattern=r"(?=(?:^\s*page\s*\d+\s*/\s*\d+\s*:|^\s*(?:protection|alerte|secours|signalisation|priorit[ée])\s*:))",
            label="sections pédagogiques FR",
            flags=re.IGNORECASE | re.MULTILINE
        )
        if chunks:
            return chunks

    # 4. Articles arabes : "المادة 1:" ou "المادة الأولى"
    if re.search(r"المادة\s+\d+", content):
        chunks = _split_by_pattern(
            doc,
            pattern=r"(?=^المادة\s+\d+)",
            label="مادة AR"
        )
        if chunks:
            return chunks

    # 5. Chapitres arabes : "=== الفصل" ou "المحور الأول"
    if re.search(r"(===\s*الفصل|المحور\s+الأول|المحور\s+الثاني)", content):
        chunks = _split_by_pattern(
            doc,
            pattern=r"(?=(?:===\s*الفصل|المحور\s+(?:الأول|الثاني|الثالث|الرابع|الخامس|السادس|السابع)))",
            label="محور AR"
        )
        if chunks:
            return chunks

    # 5.b Titres AR génériques : الفصل/الباب/القسم
    if re.search(r"^(?:الفصل|الباب|القسم)\b", content, re.MULTILINE):
        chunks = _split_by_pattern(
            doc,
            pattern=r"(?=^(?:الفصل|الباب|القسم)\b.*$)",
            label="titres AR",
            flags=re.MULTILINE
        )
        if chunks:
            return chunks

    # 5.c Titres Markdown/numérotations (utile pour docs reformattés)
    if re.search(r"^(?:#{1,3}\s+.+|\d+(?:\.\d+)*\s*[-):.])", content, re.MULTILINE):
        chunks = _split_by_pattern(
            doc,
            pattern=r"(?=^(?:#{1,3}\s+.+|\d+(?:\.\d+)*\s*[-):.]).*$)",
            label="headings génériques",
            flags=re.MULTILINE
        )
        if chunks:
            return chunks

    # 5.d Sections pedagogiques AR frequentes (accident / secours / assurance...)
    if re.search(r"(?:الحماية|الإبلاغ|الإسعاف\s+المروري|التأمين\s+وحوادث\s+المرور|أولوية\s+المرور)", content):
        chunks = _split_by_pattern(
            doc,
            pattern=r"(?=(?:^\s*(?:الحماية|الإبلاغ|الإسعاف\s+المروري|التأمين\s+وحوادث\s+المرور|أولوية\s+المرور)\s*:?|\b(?:الحماية|الإبلاغ|الإسعاف\s+المروري|التأمين\s+وحوادث\s+المرور|أولوية\s+المرور)\b))",
            label="sections pédagogiques AR",
            flags=re.MULTILINE
        )
        if chunks:
            return chunks

    # 6. Fallback : découpage classique par taille
    print(f"   ⚠️ Aucune structure détectée → découpage classique")
    return _fallback_split(doc)


def _split_by_pattern(
    doc: Document,
    pattern: str,
    label: str,
    flags: int = re.MULTILINE
) -> List[Document]:
    """
    Découpe le texte selon un pattern regex.
    Chaque section détectée devient un chunk indépendant.
    Si une section est trop longue, elle est re-découpée.
    """
    compiled = re.compile(pattern, flags)
    parts = compiled.split(doc.page_content)
    parts = [p.strip() for p in parts if p.strip()]

    if len(parts) <= 1:
        return []

    print(f"   ✅ Découpage par {label} : {len(parts)} sections")
    chunks = []

    for part in parts:
        first_line = part.split("\n")[0].strip()[:100]

        if len(part) > settings.CHUNK_SIZE * 2:
            # Section trop longue → re-découper en préservant le titre
            sub_chunks = _fallback_split(Document(
                page_content=part,
                metadata={**doc.metadata, "section": first_line}
            ))
            chunks.extend(sub_chunks)
        else:
            chunks.append(Document(
                page_content=part,
                metadata={
                    **doc.metadata,
                    "section": first_line,  # titre de la section en métadonnée
                }
            ))

    return chunks


def _fallback_split(doc: Document) -> List[Document]:
    """Découpage classique par taille avec chevauchement."""
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=settings.CHUNK_SIZE,
        chunk_overlap=settings.CHUNK_OVERLAP,
        length_function=len,
        separators=["\n\n", "\n", ".", "،", " ", ""],
    )
    return splitter.split_documents([doc])